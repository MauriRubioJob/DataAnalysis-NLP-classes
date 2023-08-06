''''


pip install python-docx


En este script de python vamos a ver cómo trabajar con un módulo que nos permite generar
archivo docx (Word, google documents)

https://python-docx.readthedocs.io/en/latest/

'''

from docx import Document

documento = Document()

# 1) tenemos que escribir todas las líneas de código que necesite mi documentos

def create_doc ( alumno , correo ):
    # Voy a diseñar esta función que me va a generar un documento distinto para cada alumno.
    documento.add_heading("Matrícula de clases", 0)

    texto = f"{alumno} con correo {correo} se ha matriculado en el curso de Python de URIA.\nCon fecha de 22/06/2020 queda satisfactoriamente matriculado."
    documento.add_paragraph(texto)

    nombre_documento = "Matrícula_"+alumno[:3]+".docx"

    documento.save(nombre_documento)




alumnos = [ ["Satur","correo_satur"] , ["Piedad" , "correo_Piedad"] , ["Alejandro" , "correo_alejandro"] , ["Nicolás", "correo_nicolas"]  ]

for alumno in alumnos:

    create_doc( alumno[0] , alumno[1])

